from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "设计说明文档.docx"
OUT = ROOT / "数据清洗过程说明文档.docx"

INK = RGBColor(30, 25, 21)
CINNABAR = RGBColor(182, 52, 43)
WOOD = RGBColor(112, 69, 54)
MUTED = RGBColor(95, 80, 71)
PAPER = "F3EBDD"
LIGHT = "FFF8EC"


def set_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    rfonts = run._element.get_or_add_rPr().rFonts
    for key in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(round(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("版上山河 · 数据清洗过程说明  ·  第 ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    tail = paragraph.add_run(" 页")
    set_font(tail, size=9, color=MUTED)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.33
    for run in p.runs:
        set_font(run, size=11, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    for run in p.runs:
        set_font(run, size=10.5, color=INK)
    if not p.runs:
        set_font(p.add_run(text), size=10.5, color=INK)
    return p


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}：")
    set_font(r, size=10.5, bold=True, color=CINNABAR)
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        shade(cell, PAPER)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            set_font(run, size=font_size, bold=True, color=WOOD)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cells[i].paragraphs[0].runs:
                set_font(run, size=font_size, color=INK)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            cells[i].paragraphs[0].paragraph_format.line_spacing = 1.1
    set_table_widths(table, widths)
    return table


doc = Document(str(TEMPLATE))
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.82)
section.bottom_margin = Inches(0.78)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.49)
section.footer_distance = Inches(0.49)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.33
for name, size, before, after in (("Heading 1", 16, 18, 10), ("Heading 2", 13, 12, 6)):
    style = doc.styles[name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = CINNABAR
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

hp = section.header.paragraphs[0]
hp.text = "《信息可视化》课程设计  |  Carved China  |  数据说明"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in hp.runs:
    set_font(run, size=8.5, color=MUTED)
add_page_number(section.footer.paragraphs[0])

# Cover: editorial cover pattern, narrative_proposal preset with a project palette override.
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
p.paragraph_format.space_after = Pt(16)
r = p.add_run("数据清洗过程说明文档")
set_font(r, size=25, bold=True, color=INK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("版上山河")
set_font(r, "SimSun", 30, True, CINNABAR)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(36)
r = p.add_run("中国传统木版年画全国对比数据集")
set_font(r, "SimSun", 16, False, WOOD)

meta = add_table(doc, ("项目", "内容"), [
    ("数据主题", "国家级非遗木版年画项目、流派、地域、历史、题材与代表色"),
    ("原始记录", "12 条首批国家级木版年画项目；14 条来源台账"),
    ("输出格式", "CSV、JSON、GeoJSON 与自动审计 JSON"),
    ("学号", "____________________________"),
    ("姓名", "____________________________"),
], [1.45, 5.05], 10.2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(34)
r = p.add_run("数据目录：data/raw  →  data/processed  →  网页可视化")
set_font(r, size=9.5, color=MUTED)
p.add_run().add_break(WD_BREAK.PAGE)

doc.add_heading("一、文档目的与数据边界", level=1)
add_body(doc, "本说明文档记录《版上山河》课程设计从资料采集到网页读取之间的数据处理过程，重点说明原始数据保存、字段标准化、人工编码、地理信息补充、质量审计和输出文件之间的关系。其目的不是把所有处理包装成完全自动化流程，而是明确区分官方事实、资料摘录、设计编码与近似推断，使数据可追溯、可复核、可重复使用。")
add_body(doc, "项目的统计范围限定为中国非物质文化遗产网第一批国家级非物质文化遗产代表性项目名录中的木版年画项目，共 12 条，覆盖 11 个省级地区。该范围适合进行全国代表性流派比较，但不能被理解为中国全部木版年画产地或全部作品的普查。历史、题材、工艺与色彩信息来自项目详情页、地方政府和文化机构公开资料；其中题材分类、代表色和起源排序年份属于面向可视化的二次编码。")
add_note(doc, "口径提示", "名录字段是官方事实；originYear 是排序辅助值；subjects 与 colors 是设计编码。网页及报告不得将后两者表述为官方统计或全量图像测量结果。")

doc.add_heading("二、文件与数据流", level=1)
add_body(doc, "数据文件采用“原始层—处理层—展示层”三级组织。原始层保留采集时的中文字段、来源 URL 和访问日期；处理层建立统一英文键名并补充可视化所需字段；展示层由 ECharts 和 p5.js 读取 JSON 或 GeoJSON，不直接修改数据。来源台账独立保存，避免出处散落在代码注释中。")
add_table(doc, ("阶段", "文件", "作用"), [
    ("原始层", "data/raw/国家级木版年画首批名录.csv", "保留项目序号、编号、名称、地区、保护单位和来源"),
    ("来源层", "data/sources.csv", "记录 14 个资料来源、发布者、用途和访问日期"),
    ("处理层", "data/processed/projects.json", "供地图、时间轴、题材图、色彩图和生成艺术读取"),
    ("地图层", "data/processed/china.geojson", "提供省级行政区边界，仅作为地图底图"),
    ("审计层", "data/processed/audit-summary.json", "输出数量、覆盖率、频次和错误列表"),
], [1.0, 2.65, 2.85])

doc.add_heading("三、清洗与编码步骤", level=1)
steps = [
    ("1. 原始数据固化", "从国家级名录提取项目序号、项目编号、名称、类别、公布时间、申报地区、保护单位和来源链接。CSV 使用 UTF-8 编码；含逗号的保护单位以双引号包裹，并记录访问日期。"),
    ("2. 名称和地区标准化", "统一省级地区写法，将县、市、区分离为 province 与 city。重庆、天津按省级行政区处理；坐标使用十进制度数，longitude 与 latitude 均为数值。"),
    ("3. 稳定标识建立", "为每条项目建立小写英文 id，例如 yangliuqing、wuqiang、taohuawu。id 用于筛选、联动与测试，不随中文展示名称变化；projectNo 与 code 保留官方名录标识。"),
    ("4. 历史字段归一", "origin 保留“北宋”“宋末元初”“明代成化年间”等原始文字表述；originYear 仅取近似年份用于时间轴排序。近似值不替代历史断代，也不在界面中伪装为精确创立年份。"),
    ("5. 题材分类编码", "把公开资料中的题材描述归并为门神驱邪、吉祥祈福、娃娃仕女、戏曲故事、花鸟瑞兽、民俗生活六类。一条流派可以对应多个题材，使用字符串数组保存。"),
    ("6. 色彩字段编码", "将朱砂、藤黄、石绿和墨色统一为 #B6342B、#D5A33A、#39705A、#1E1915。HEX 转为大写并验证格式；这些值表达视觉倾向，不代表对全部馆藏作品逐像素测色。"),
    ("7. 工艺与来源补充", "technique 摘要保留各流派主要工艺差异；sourceUrl 指向名录，detailSource 指向项目详情或地方权威资料，accessed 记录访问日期。"),
    ("8. 自动审计与输出", "运行 node tools/audit-data.mjs，对重复 id、缺失来源、无效坐标、未知题材和无效 HEX 进行检查，并生成 audit-summary.json。错误列表非空时命令以失败状态结束。"),
]
for title, detail in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_font(r, size=11.2, bold=True, color=WOOD)
    add_body(doc, detail)

doc.add_heading("四、字段映射与转换规则", level=1)
add_table(doc, ("原始/资料字段", "处理字段", "类型", "清洗规则"), [
    ("项目序号", "projectNo", "Number", "转为整数；用于核对名录顺序"),
    ("项目编号", "code", "String", "保留官方编号文字"),
    ("项目名称", "school", "String", "去除首尾空白，保留正式名称"),
    ("申报地区", "province / city", "String", "拆分省级与城市/区县层级"),
    ("起源描述", "origin / originYear", "String / Number", "原文保留；另建近似排序年份"),
    ("题材描述", "subjects", "Array", "映射到六类受控词表，可多选"),
    ("传统用色", "colors", "Array", "转换为大写六位 HEX 数组"),
    ("来源", "sourceUrl / detailSource", "URL", "名录与细节来源分开保存"),
], [1.42, 1.62, 1.0, 2.46], 8.8)
add_note(doc, "示例", "“明末清初”保留在 origin 中，同时设置 originYear=1640 仅用于排序；“朱砂、藤黄、石绿”转为三个 HEX 值，便于图表和生成艺术共享同一色板。")

doc.add_heading("五、质量控制与审计结果", level=1)
add_body(doc, "审计脚本由 tools/audit-data.mjs 调用 scripts/data-model.mjs。validateProjects 首先验证 id 唯一性、来源存在性和 HEX 合法性；随后检查 detailSource、经纬度、题材受控词表，并汇总项目、地区、流派、题材和颜色频次。当前审计结果 errors 为空，说明所有强制规则均通过。")
add_table(doc, ("检查项", "结果", "判定"), [
    ("项目记录", "12 条", "与首批名录样本一致"),
    ("省级地区", "11 个", "山东包含 2 个项目，其余各 1 个"),
    ("详情来源覆盖率", "100%", "12/12 均有 detailSource"),
    ("坐标覆盖率", "100%", "12/12 均有有效经纬度"),
    ("题材覆盖率", "100%", "12/12 均至少有一个题材"),
    ("代表色覆盖率", "100%", "12/12 均至少有一个 HEX 颜色"),
    ("规则错误", "0 项", "重复、缺失、未知枚举与格式检查通过"),
], [2.05, 1.25, 3.2], 9.3)
add_body(doc, "频次汇总用于网页图表：吉祥祈福 11 项、戏曲故事 9 项、门神驱邪 8 项、民俗生活 6 项、花鸟瑞兽 4 项、娃娃仕女 3 项；朱砂色覆盖 12 项、藤黄色 11 项、石绿色 10 项、墨色 6 项。频次表示编码字段在 12 个流派中的出现数量，不是作品件数。")

doc.add_heading("六、复现方法", level=1)
add_body(doc, "项目根目录不需要数据库。安装 Node.js 后可在 PowerShell 或命令提示符中运行审计命令；也可直接检查已提交的 audit-summary.json。网页通过本地 HTTP 服务读取数据，不能直接双击 index.html 依赖 file 协议加载。")
add_table(doc, ("操作", "命令/入口"), [
    ("运行数据审计", "node tools/audit-data.mjs"),
    ("运行全部测试", "node --test tests/*.test.mjs"),
    ("本地启动", "双击 启动项目.cmd，Microsoft Edge 自动打开"),
    ("在线检查", "https://hans200406.github.io/Carved-China/"),
], [1.55, 4.95], 9.5)
add_bullet(doc, "修改 projects.json 后必须重新执行数据审计。")
add_bullet(doc, "新增题材时应同时更新 ALLOWED_SUBJECTS 受控词表和图表标签。")
add_bullet(doc, "新增颜色时必须使用六位 HEX，并确认生成艺术模块能够读取。")
add_bullet(doc, "修改坐标后应在地图中检查标记是否落在对应省级区域。")

doc.add_heading("七、局限性与后续扩展", level=1)
add_body(doc, "当前数据集规模较小，强调课程设计中的方法完整性与可追溯性。坐标用于项目所在地标记，并非保护单位建筑的测绘坐标；题材与色彩为资料描述基础上的人工编码，可能受资料详略影响；originYear 是时间轴排序的近似值。后续若扩展第二批及扩展项目，应保留批次字段，并为不同来源的同名项目建立去重规则。")
add_body(doc, "进一步扩展可引入博物馆开放图像，对代表作品做颜色聚类和纹样识别；也可以加入传承人、作坊、销售网络与传播路线，形成多实体关系数据。在任何扩展中，原始事实、人工编码和算法推断都应使用不同字段保存，并继续维护来源台账与自动审计。")

doc.add_heading("附录：提交文件核对", level=1)
for item in (
    "原始名录 CSV 已保留来源 URL 与访问日期。",
    "清洗后 JSON 已包含地图、时间轴、题材、色彩与生成艺术所需字段。",
    "来源台账包含 14 条来源及其具体用途。",
    "审计 JSON 错误列表为空，四项覆盖率均为 100%。",
    "本说明文档中的学号、姓名保持空白，提交前由本人填写。",
):
    add_bullet(doc, item)

doc.core_properties.title = "版上山河：中国传统木版年画数据清洗过程说明"
doc.core_properties.subject = "《信息可视化》课程设计数据说明"
doc.core_properties.author = ""
doc.core_properties.keywords = "木版年画; 数据清洗; 信息可视化; 数据审计"
doc.save(str(OUT))
print(OUT)
