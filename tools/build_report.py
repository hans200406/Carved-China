from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"E:\作业\信息可视化设计\23120032038")
TEMPLATE = Path(r"E:\作业\信息可视化设计\23120032037-蔡绍涵-非遗课设\设计说明文档.docx")
OUT = ROOT / "设计说明文档.docx"
ASSETS = ROOT / "docs" / "report-assets"

INK = RGBColor(30, 25, 21)
CINNABAR = RGBColor(182, 52, 43)
WOOD = RGBColor(112, 69, 54)
MUTED = RGBColor(95, 80, 71)
PAPER = "F3EBDD"
LIGHT = "FFF8EC"


def set_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(row.cells[idx])


def add_page_number(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("版上山河：中国传统木版年画多维信息可视化设计  ·  第 ")
    set_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    run2 = paragraph.add_run(" 页")
    set_font(run2, size=9, color=MUTED)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        set_font(r, size=11, color=INK)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_font(r, size=9, color=MUTED)
    return p


def add_figure(doc, filename, caption, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(ASSETS / filename), width=Inches(width))
    add_caption(doc, caption)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)


def add_reference(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.34)
    p.paragraph_format.first_line_indent = Inches(-0.17)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_font(r, size=9, color=INK)


doc = Document(str(TEMPLATE))
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.5

for name, size, before, after in (("Heading 1", 16, 18, 10), ("Heading 2", 13, 12, 6)):
    style = doc.styles[name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = CINNABAR
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.text = "《信息可视化》课程设计说明书  |  Carved China"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    set_font(r, size=8.5, color=MUTED)
add_page_number(section.footer.paragraphs[0])

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(18)
r = p.add_run("《信息可视化》课程设计说明书")
set_font(r, "Microsoft YaHei", 23, True, INK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("版上山河")
set_font(r, "SimSun", 32, True, CINNABAR)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
r = p.add_run("中国传统木版年画多维信息可视化设计")
set_font(r, "SimSun", 17, False, WOOD)

meta = doc.add_table(rows=6, cols=2)
meta.style = "Table Grid"
set_table_widths(meta, [1.55, 4.55])
cover_rows = [
    ("课程名称", "信息可视化"),
    ("课设主题", "数字遗珍——非遗文化多维信息可视化设计"),
    ("作品形式", "交互网页 / H5 数字媒体项目"),
    ("选题对象", "中国传统木版年画"),
    ("学号", "____________________________"),
    ("姓名", "____________________________"),
]
for idx, (label, value) in enumerate(cover_rows):
    c0, c1 = meta.rows[idx].cells
    c0.text, c1.text = label, value
    shade(c0, PAPER)
    c0.vertical_alignment = c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for r in c0.paragraphs[0].runs:
        set_font(r, size=10.5, bold=True, color=WOOD)
    for r in c1.paragraphs[0].runs:
        set_font(r, size=10.5, color=INK)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(36)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("在线作品：https://hans200406.github.io/Carved-China/")
set_font(r, size=9.5, color=MUTED)
p.add_run().add_break(WD_BREAK.PAGE)

# Chapter 1
doc.add_heading("一、项目背景与选题意义", level=1)
add_body(doc, "本课程设计以中国传统木版年画为研究与创作对象，作品中文名为“版上山河”，英文名为“Carved China”。木版年画是伴随春节礼俗、雕版印刷和民间信仰长期发展起来的传统视觉艺术。它既是门神、吉祥娃娃、戏曲故事、花鸟瑞兽等图像的载体，也是观察地方城镇、手工业网络、审美习惯和生活愿望的重要文化样本。不同产地在造型、线条、工序和配色上各具特点，例如杨柳青年画强调木版套印与手工彩绘结合，桃花坞年画具有江南印刷传统，朱仙镇与武强等北方流派线条朴厚，绵竹、滩头等西南流派又形成独特的手绘或本地造纸工序。")
add_body(doc, "课程任务要求同时完成理性的数据分析和感性的图形艺术两个模块，并将二者整合为交互数字媒体作品。因此，本项目没有将年画处理为静态图片展览，而是建立“全国流派宏观比较—历史与题材结构分析—传统色彩编码—生成艺术重构”的叙事链条。模块A使用真实名录与权威资料建立数据基础，模块B把门神、莲花、鲤鱼、花鸟和云纹等元素转化为可计算的几何图元，模块C通过滚动、悬停、点击、筛选和参数调节将两部分连接起来。")
add_body(doc, "选择全国木版年画进行对比具有三方面意义。第一，年画项目分布跨越华北、华东、华中、华南和西南，适合通过地图呈现地域谱系。第二，各流派的形成时代、题材功能和制作工艺差异显著，可支持时间轴、题材频次和色彩比较。第三，年画本身具有对称、重复、套色和符号化特征，适合转译为p5.js生成艺术。通过信息可视化，观众不仅能“看到年画”，还可以理解不同流派为何形成差异，并通过交互参与新的数字年画生成过程。")
add_figure(doc, "00-process.png", "图1  课程设计流程草图：资料采集、清洗编码、模块A、模块B与模块C的关系")

# Chapter 2
doc.add_heading("二、数据来源与处理方法", level=1)
add_body(doc, "项目以中国非物质文化遗产网的国家级非物质文化遗产代表性项目名录为主来源，整理第一批传统美术类木版年画项目。基础字段包括项目序号、项目编号、项目名称、公布时间、类型、申报地区和保护单位；历史、题材与工艺信息则由中国非遗网项目详情页、地方政府、文化和旅游部门及博物馆公开资料补充。最终选取杨柳青、武强、桃花坞、漳州、杨家埠、高密扑灰、朱仙镇、滩头、佛山、梁平、绵竹和凤翔共12项样本，覆盖11个省级地区。每条记录均保存来源链接与访问日期，便于追溯。")
add_body(doc, "数据处理遵循“原始记录保留、标准字段另建、推断内容明确标注”的原则。首先统一省级地区和城市名称，并补充用于地图表达的近似经纬度；其次把“宋末元初”“明代成化年间”“明末清初”等文字年代归一为用于排序的近似年份，同时保留原始历史表述。这个年份只表示流派形成阶段，不代表精确创立日期。再次，将公开资料中的题材描述归并为门神驱邪、吉祥祈福、娃娃仕女、戏曲故事、花鸟瑞兽、民俗生活六个类别，以支持跨流派比较。")
add_body(doc, "色彩字段属于设计编码，而不是官方作品像素统计。项目根据权威资料对代表作品和传统用色的描述，将朱砂、藤黄、石绿、墨色与木褐整理为统一HEX值，用于网页图表和生成艺术色板。为防止把设计判断误写成客观统计，网页与报告均明确提示：色彩比较表达视觉倾向，不代表对全部馆藏作品进行测色。项目还编写了自动数据审计脚本，对重复ID、来源缺失、坐标缺失、HEX格式和题材枚举进行检查。当前12条记录在细节来源、坐标、题材和色彩字段上的覆盖率均为100%。")

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
set_table_widths(table, [1.45, 2.3, 2.35])
for i, v in enumerate(("字段", "含义", "用途")):
    table.rows[0].cells[i].text = v
    shade(table.rows[0].cells[i], PAPER)
fields = [
    ("项目名称/编号", "国家级项目名称及Ⅶ类编号", "识别项目与名录依据"),
    ("省市/坐标", "申报地区及近似经纬度", "地图分布与空间比较"),
    ("起源表述/年份", "原始历史描述及归一化排序值", "历史时间轴"),
    ("题材类别", "六类统一的内容编码", "跨流派题材比较"),
    ("代表色", "传统色名及HEX设计编码", "色彩比较与生成艺术"),
    ("来源链接", "名录与项目细节页面", "数据追溯与核验"),
]
for row in fields:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in cells[i].paragraphs[0].runs:
            set_font(r, size=9.5, color=INK)
for c in table.rows[0].cells:
    for r in c.paragraphs[0].runs:
        set_font(r, size=9.5, bold=True, color=WOOD)
add_caption(doc, "表1  木版年画数据字段及其可视化用途")

# Chapter 3
doc.add_heading("三、模块A：数据分析信息可视化设计", level=1)
add_body(doc, "模块A围绕“木版年画如何在中国形成多中心流派”建立宏观画像。项目采用全国地图、历史时间轴、题材频次图和传统色彩比较图四种形式。全国地图使用本地GeoJSON边界与项目经纬度标记，既能显示跨区域分布，又允许用户缩放和点击项目点。地图点击某一流派后，题材图和色彩图会同步筛选，从而避免地图成为孤立装饰。")
add_body(doc, "历史时间轴使用散点而非连续折线，因为各流派的形成时间不是同一变量的连续观测值，也不存在可直接相连的数值趋势。纵轴列出流派，横轴显示归一后的年代，悬停时呈现“北宋”“宋末元初”“明代成化年间”等原始表述。题材频次采用横向条形图，便于容纳中文类别名称并进行数量排序；色彩比较使用竖向柱状图，将朱砂、藤黄、石绿和墨色的流派覆盖数量与对应色块结合，使数值和视觉编码保持一致。")
add_figure(doc, "01-map.png", "图2  模块A全国木版年画项目分布地图", width=6.05)
add_figure(doc, "02-charts.png", "图3  模块A历史演进时间轴", width=6.05)
add_body(doc, "从当前样本中可以提炼三条主要洞察。第一，12项首批国家级样本分布在11个省级地区，形成明显的多中心格局；它们既包含天津、苏州、开封、潍坊等商业与交通节点，也包含武强、滩头、梁平、绵竹等具有地方手工业基础的县域或城镇，说明年画传播与节庆市场、纸张印刷条件和地方民俗生态密切相关。第二，在统一题材编码中，吉祥祈福出现于11项，戏曲故事出现于9项，门神驱邪出现于8项，表明年画首先服务于春节空间中的祈愿、镇宅和叙事功能，同时又吸收地方戏曲、历史传说与日常生活。第三，朱砂出现在全部12项代表色编码中，藤黄与石绿也具有较高覆盖率；共同色形成“年味”的视觉共识，而套印顺序、手绘方式和色彩比例又构成流派差异。")
add_body(doc, "上述结论的边界同样需要说明：样本范围是第一批国家级木版年画项目，并非中国全部年画产地；题材与色彩来自公开资料的描述性编码，不能替代大规模作品图像统计。因此，网页在图表旁保留数据口径提示，并在报告中公开来源与处理原则。这样的限制说明能够避免过度推断，也使后续扩展第二批及扩展项目时有清晰方法可循。")

# Chapter 4
doc.add_heading("四、模块B：图形艺术信息可视化设计", level=1)
add_body(doc, "模块B围绕年画的微观美学展开。课程任务要求不能直接堆砌照片，因此本项目没有把网络图片裁切成拼贴，而是以代码生成艺术重构年画的对称、重复、边框和套色规律。图形系统将莲花、盾形门神结构、鲤鱼轮廓、花瓣、菱形纹样与边框转化为Canvas几何图元，并通过旋转、缩放、环形排列和随机种子组成新的数字年画。")
add_body(doc, "生成器提供流派、题材、密度和随机种子四类控制。选择流派会切换经设计整理的代表色板；选择“吉祥祈福”“门神驱邪”或“花鸟瑞兽”会改变中心符号与放射节奏；密度滑杆控制纹样环数与复杂度；“重新套印”按钮改变随机种子，使图元位置和尺寸产生变化。系统保持确定性随机逻辑，同一参数和种子能够重现相同结果。用户还可以导出当前画布为PNG，用于保存个人生成作品。")
add_body(doc, "模块B与模块A之间通过数据字段建立联系。模块A中的题材编码决定生成器可选择的主题，流派色彩编码决定画布用色，历史与地域分析则为不同视觉风格提供文化背景。这样，生成艺术不是与数据无关的装饰，而是数据分析结果向视觉语言的转译。")
add_figure(doc, "03-studio.png", "图4  模块B生成艺术与色彩参数界面", width=5.35)

# Chapter 5
doc.add_heading("五、模块C：文化叙事与交互设计", level=1)
add_body(doc, "模块C负责把理性分析与感性生成整合为连续体验。页面采用“开场吸引—全国分布—百年流变—题材功能—传统色彩—生成艺术—保护结语”的纵向叙事顺序。开场以宣纸暖白、朱砂和巨大的“版”字残影建立文化氛围；中段交替使用暖白、深墨和朱砂章节，让数据阅读保持节奏；生成艺术区转为深色背景，突出用户操作和套色画布；结尾以三条洞察回收全文。")
add_body(doc, "交互设计包含三个层面。第一，ECharts图表支持悬停提示、地图缩放和项目点击，地图选择会联动题材与色彩。第二，固定导航允许用户跳转到“年画中国、百年流变、图必有意、五色成章、版上生花”等章节，滚动进入时通过IntersectionObserver触发淡入和位移动画。第三，p5.js画布提供下拉框、滑杆、按钮与图片导出。手机端不依赖悬停，控件触控区域不小于44像素；系统检测prefers-reduced-motion，在用户要求减少动画时直接显示内容。")
add_body(doc, "视觉规范参考暖色编辑式网页的留白、字体分工和内容节奏，但重新建立了木版年画专用设计令牌。页面以宣纸暖白为画布，墨色为正文，朱砂为主强调，藤黄与石绿作为数据辅助色；标题使用宋体类字体，正文与控件使用黑体类字体。卡片只使用细边框、轻微纸张色差和克制圆角，避免霓虹、玻璃拟态和无意义渐变，使传统文化气息与现代数据界面保持平衡。")

# Chapter 6
doc.add_heading("六、作品实现与文件结构", level=1)
add_body(doc, "作品采用纯静态前端架构，不依赖Vue、React、后端数据库或复杂构建链，便于课程验收和GitHub Pages部署。ECharts与p5.js均保存在vendor目录，断网时仍可在本地服务器运行。数据以JSON和CSV保存，地图使用本地中国省级行政区GeoJSON。项目提供一键启动脚本，双击后启动本地服务器并尝试在Microsoft Edge中打开；在线版本由GitHub Actions自动发布。")
add_bullet(doc, "index.html：七个叙事章节、导航、图表容器和生成艺术控制。")
add_bullet(doc, "styles/：设计令牌、响应式布局、章节配色与减少动画规则。")
add_bullet(doc, "scripts/app.mjs：数据加载、中国地图、时间轴、题材与色彩图表、筛选联动和滚动叙事。")
add_bullet(doc, "scripts/art-model.mjs：题材中心符号、流派色板和生成艺术模型。")
add_bullet(doc, "data/raw 与 data/processed：原始名录、结构化项目数据、地图和审计结果。")
add_bullet(doc, "tools/audit-data.mjs：来源、坐标、题材、颜色与重复项自动审计。")
add_bullet(doc, "poster/：1080×1920展示海报PNG及HTML源文件。")
add_bullet(doc, "tests/：数据模型、生成艺术、页面结构、海报与部署共27项自动测试。")
add_body(doc, "项目已发布至 https://hans200406.github.io/Carved-China/ 。线上验证包括12项数据加载、中国地图渲染、流派筛选、生成画布、手机端横向溢出检查和资源路径检查。GitHub Pages工作流在master分支更新后自动部署，保证提交版本与在线版本保持一致。")

# Chapter 7
doc.add_heading("七、总结与反思", level=1)
add_body(doc, "“版上山河”尝试把中国传统木版年画从静态观赏对象转化为可分析、可交互、可再生成的数字文化体验。模块A以真实名录与权威资料呈现12项首批国家级项目的地域、历史、题材和色彩结构；模块B把莲花、门神、鲤鱼和传统色转化为可控制的生成规则；模块C再通过滚动叙事与联动交互把宏观数据和微观美学连接起来。作品既回应了课程对数据严谨性的要求，也避免直接堆砌照片，强调图形重构与信息传达。")
add_body(doc, "项目的核心价值在于将“看见传统”和“理解传统”结合。地图与时间轴帮助观众建立全国流派框架，题材和色彩比较揭示共性与差异，生成艺术则让用户通过参数操作体会对称、套色与重复的造型规律。设计过程中最重要的反思是区分事实、编码与推断：名录字段属于官方事实，年代归一属于排序方法，题材与色彩属于设计编码。只有公开这些层次，信息可视化才不会因视觉效果而牺牲准确性。")
add_body(doc, "后续可进一步扩展第二批及扩展项目，引入更多博物馆开放图像，对代表作品进行颜色聚类和纹样识别；也可以增加传承人、作坊和传播路线数据，构建更完整的师承与流通网络。当前版本已经完成网页、数据集、生成艺术源文件、在线发布和展览海报，为继续拓展奠定了可复用的数据结构和视觉系统。")

doc.add_heading("参考资料", level=1)
refs = [
    "中国非物质文化遗产网：国家级非物质文化遗产代表性项目名录，https://www.ihchina.cn/project.html?tid=7。",
    "中国非物质文化遗产网：武强、高密扑灰、朱仙镇、滩头、杨家埠、绵竹等项目详情页。",
    "福建省人民政府：漳州木版年画相关资料。",
    "重庆市文化和旅游发展委员会：梁平木版年画资料。",
    "宝鸡市人民政府：凤翔木版年画资料。",
    "广东省文化和旅游厅、中国非遗网：佛山木版年画工艺与传承资料。",
    "DataV GeoAtlas：中华人民共和国省级行政区边界GeoJSON。",
]
for ref in refs:
    add_reference(doc, ref)

doc.core_properties.title = "版上山河：中国传统木版年画多维信息可视化设计说明书"
doc.core_properties.subject = "《信息可视化》课程设计"
doc.core_properties.author = ""
doc.core_properties.keywords = "木版年画; 信息可视化; ECharts; p5.js; 非物质文化遗产"
doc.save(str(OUT))
print(OUT)
